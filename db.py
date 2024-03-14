from docx import Document
from main import collect_messages, insert_answers_into_doc, find_start, find_end
import docx_tools
import main

doc = Document()
doc.add_paragraph("<gpt>")
doc.add_paragraph("</gpt>")
answers = main.collect_messages(doc)
main.insert_answers_into_doc(answers, doc)
assert docx_tools.doc_text(doc) == ""
# AssertionError: assert 'FOO>BAR>' == 'FOOBAR'