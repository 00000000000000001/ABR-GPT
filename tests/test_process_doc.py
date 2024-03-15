import sys

sys.path.append("../")
sys.path.append("./")
from main import collect_messages, insert_answers_into_doc, find_end, find_start
from docx import Document
import docx_tools


def test_abc():
    doc = Document()
    p = doc.add_paragraph("<gpt>abc</gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    insert_answers_into_doc(messages, doc)
    assert find_start(p) == -1
    assert find_end(p) == -1
    assert docx_tools.combineDocText(doc) == "TEST"


def test_empty():
    doc = Document()
    p = doc.add_paragraph("<gpt></gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    insert_answers_into_doc(messages, doc)
    assert find_start(p) == -1
    assert find_end(p) == -1
    assert docx_tools.combineDocText(doc) == "TEST"


def test_empty_two_tags():
    doc = Document()
    p = doc.add_paragraph("<gpt></gpt><gpt></gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    messages[1][0] = "BAUM"
    insert_answers_into_doc(messages, doc)
    assert find_start(p) == -1
    assert find_end(p) == -1
    assert docx_tools.combineDocText(doc) == "TESTBAUM"


def test_empty_two_paragraphs():
    doc = Document()
    p1 = doc.add_paragraph("<gpt>")
    p2 = doc.add_paragraph("</gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    insert_answers_into_doc(messages, doc)
    assert find_start(p1) == -1
    assert find_end(p1) == -1
    assert find_start(p2) == -1
    assert find_end(p2) == -1
    assert docx_tools.combineDocText(doc) == "TEST"


def test_empty_two_paragraphs_two_tags():
    doc = Document()
    p1 = doc.add_paragraph("<gpt>")
    p2 = doc.add_paragraph("</gpt><gpt></gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    messages[1][0] = "BAUM"
    insert_answers_into_doc(messages, doc)
    assert find_start(p1) == -1
    assert find_end(p1) == -1
    assert find_start(p2) == -1
    assert find_end(p2) == -1
    assert docx_tools.combineDocText(doc) == "TESTBAUM"


def test_nonempty_two_paragraphs_two_tags1():
    doc = Document()
    p1 = doc.add_paragraph("<gpt>")
    p2 = doc.add_paragraph("FOO</gpt><gpt>BAR</gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    messages[1][0] = "BAUM"
    insert_answers_into_doc(messages, doc)
    assert find_start(p1) == -1
    assert find_end(p1) == -1
    assert find_start(p2) == -1
    assert find_end(p2) == -1
    assert docx_tools.combineDocText(doc) == "TESTBAUM"


def test_nonempty_two_paragraphs_two_tags2():
    doc = Document()
    p1 = doc.add_paragraph("<gpt>FOO")
    p2 = doc.add_paragraph("</gpt><gpt>BAR</gpt>")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    messages[1][0] = "BAUM"
    insert_answers_into_doc(messages, doc)
    assert find_start(p1) == -1
    assert find_end(p1) == -1
    assert find_start(p2) == -1
    assert find_end(p2) == -1
    assert docx_tools.combineDocText(doc) == "TESTBAUM"


def test_nonempty_two_paragraphs_two_tags_and_text_inbetween():
    doc = Document()
    p1 = doc.add_paragraph("a<gpt>FOO")
    p2 = doc.add_paragraph("</gpt>b<gpt>BAR</gpt>c")
    messages = collect_messages(doc)
    messages[0][0] = "TEST"
    messages[1][0] = "BAUM"
    insert_answers_into_doc(messages, doc)
    assert find_start(p1) == -1
    assert find_end(p1) == -1
    assert find_start(p2) == -1
    assert find_end(p2) == -1
    assert docx_tools.combineDocText(doc) == "aTESTbBAUMc"
