import sys
sys.path.append('../')
sys.path.append('./')
import main
from docx import Document

def test_collect_messages1():
    doc = Document()
    doc.add_paragraph("")
    assert main.collect_messages(doc) == []


def test_collect_messages2():
    doc = Document()
    doc.add_paragraph("<gpt>")
    assert main.collect_messages(doc) == []


def test_collect_messages_empty():
    doc = Document()
    doc.add_paragraph("<gpt></gpt>")
    assert main.collect_messages(doc) == [["", 5, 0, 4, 0]]


def test_collect_messages_nonempty():
    doc = Document()
    doc.add_paragraph("<gpt>ASD</gpt>")
    assert main.collect_messages(doc) == [["ASD", 5, 0, 7, 0]]


def test_collect_multiple_messages_nonempty():
    doc = Document()
    doc.add_paragraph("<gpt>ASD</gpt><gpt>QWE</gpt>")
    assert main.collect_messages(doc) == [["ASD", 5, 0, 7, 0], ["QWE", 19, 0, 21, 0]]


def test_collect_messages_is_empty_on_multiple_paragraphs1():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("")
    assert main.collect_messages(doc) == []


def test_collect_messages_is_empty_on_multiple_paragraphs2():
    doc = Document()
    doc.add_paragraph("FOO")
    doc.add_paragraph("BAR")
    assert main.collect_messages(doc) == []


def test_collect_messages_is_empty_on_multiple_paragraphs3():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("BAR")
    assert main.collect_messages(doc) == []


def test_collect_messages_is_empty_on_multiple_paragraphs4():
    doc = Document()
    doc.add_paragraph("<gpt></gpt>")
    doc.add_paragraph("BAR")
    assert main.collect_messages(doc) == [["", 5, 0, 4, 0]]


def test_collect_messages_is_not_empty_on_multiple_paragraphs1():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("</gpt>BAR")
    assert main.collect_messages(doc) == [["", 5, 0, -1, 1]]


def test_collect_messages_is_not_empty_on_multiple_paragraphs2():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("BAR</gpt>")
    assert main.collect_messages(doc) == [["BAR", 5, 0, 2, 1]]


def test_collect_messages_is_not_empty_on_multiple_paragraphs3():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("BAR")
    doc.add_paragraph("</gpt>")
    assert main.collect_messages(doc) == [["BAR", 5, 0, -1, 2]]


def test_collect_messages_is_not_empty_on_multiple_paragraphs4():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("BAR")
    doc.add_paragraph("FOO")
    doc.add_paragraph("</gpt>")
    assert main.collect_messages(doc) == [["BARFOO", 5, 0, -1, 3]]


def test_collect_messages_is_not_empty_on_multiple_paragraphs5():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("FOO")
    doc.add_paragraph("</gpt>")
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("BAR")
    doc.add_paragraph("</gpt>")
    assert main.collect_messages(doc) == [["FOO", 5, 0, -1, 2], ["BAR", 5, 3, -1, 5]]


def test_collect_messages_is_empty_on_multiple_paragraphs():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("</gpt>")
    assert main.collect_messages(doc) == [["", 5, 0, -1, 1]]