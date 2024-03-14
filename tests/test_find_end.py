import sys
sys.path.append('../')
sys.path.append('./')
import main
from docx import Document

def test_find_end_on_empry_string_is_minus_one():
    doc = Document()
    p = doc.add_paragraph("")
    assert main.find_end(p) == -1


def test_find_end_is_zero():
    doc = Document()
    p = doc.add_paragraph("</gpt>")
    assert main.find_end(p) == 0


def test_find_end_is_three():
    doc = Document()
    p = doc.add_paragraph("FOO</gpt>")
    assert main.find_end(p) == 3


def test_find_end_finds_first():
    doc = Document()
    p = doc.add_paragraph("FOO</gpt></gpt>")
    assert main.find_end(p) == 3