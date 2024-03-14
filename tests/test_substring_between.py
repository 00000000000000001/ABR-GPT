import sys
sys.path.append('../')
sys.path.append('./')
import main
from docx import Document
import pytest

def test_substring_between_in_empty_string_returns_string():
    doc = Document()
    doc.add_paragraph("")
    assert main.substring_between(0, 0, 5, 0, doc) == ""


def test_substring_between_in_empty_string_returns_empty_string1():
    doc = Document()
    doc.add_paragraph("")
    assert main.substring_between(0, 0, 0, 0, doc) == ""


def test_substring_between_in_empty_string_returns_empty_string2():
    doc = Document()
    doc.add_paragraph("")
    assert main.substring_between(0, 0, 1, 0, doc) == ""


def test_substring_between_in_empty_string_returns_nonempty_string1():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(0, 0, 1, 0, doc) == "a"


def test_substring_between_in_empty_string_returns_nonempty_string2():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(0, 0, 2, 0, doc) == "as"


def test_substring_between_in_empty_string_returns_nonempty_string3():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(0, 0, 3, 0, doc) == "asd"


def test_substring_between_in_empty_string_returns_nonempty_string4():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(0, 0, 5, 0, doc) == "asd"


def test_substring_between_in_empty_string_returns_nonempty_string5():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(1, 0, 3, 0, doc) == "sd"


def test_substring_between_in_empty_string_returns_nonempty_string6():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(2, 0, 3, 0, doc) == "d"


def test_substring_between_in_empty_string_returns_nonempty_string7():
    doc = Document()
    doc.add_paragraph("asd")
    assert main.substring_between(3, 0, 3, 0, doc) == ""


def test_substring_betweeen_to_empty_paragraphs_is_empty():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("")
    assert main.substring_between(0, 0, 0, 1, doc) == ""


def test_substring_betweeen_to_nonempty_paragraphs_is_empty():
    doc = Document()
    doc.add_paragraph("AUTO")
    doc.add_paragraph("BAUM")
    assert main.substring_between(0, 1, 0, 1, doc) == ""

def test_start_paragraph_is_out_of_bounds1():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        main.substring_between(0, -1, 0, 0, doc)


def test_start_paragraph_is_out_of_bounds2():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        main.substring_between(0, 3, 0, 0, doc)


def test_start_paragraph_is_out_of_bounds3():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        main.substring_between(0, 0, 0, -1, doc)


def test_start_paragraph_is_out_of_bounds4():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        main.substring_between(0, 0, 0, 3, doc)


def test_substring_betweeen_to_nonempty_paragraphs_is_nonempty1():
    doc = Document()
    doc.add_paragraph("AUTO")
    doc.add_paragraph("BAUM")
    assert main.substring_between(0, 1, 4, 1, doc) == "BAUM"

def test_substring_betweeen_to_nonempty_paragraphs_is_nonempty2():
    doc = Document()
    doc.add_paragraph("AUTO")
    doc.add_paragraph("BAUM")
    assert main.substring_between(0, 0, 0, 1, doc) == "AUTO"