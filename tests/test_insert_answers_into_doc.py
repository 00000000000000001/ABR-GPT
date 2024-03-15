from docx import Document
import sys
sys.path.append('../')
sys.path.append('./')
import docx_tools
import main


def test_insert_answers_into_doc1():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("FOO")
    doc.add_paragraph("</gpt>")
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("BAR")
    doc.add_paragraph("</gpt>")
    answers = main.collect_messages(doc)
    main.insert_answers_into_doc(answers, doc)
    assert docx_tools.combineDocText(doc) == "FOOBAR"


def test_insert_answers_into_doc2():
    doc = Document()
    doc.add_paragraph("<gpt>")
    doc.add_paragraph("</gpt>")
    answers = main.collect_messages(doc)
    main.insert_answers_into_doc(answers, doc)
    assert docx_tools.combineDocText(doc) == ""


def test_insert_answers_into_doc3():
    doc = Document()
    doc.add_paragraph("")
    answers = main.collect_messages(doc)
    main.insert_answers_into_doc(answers, doc)
    assert docx_tools.combineDocText(doc) == ""