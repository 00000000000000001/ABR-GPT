import sys
import docx_tools
import docx
import gpt
import fcntl
import config
import glob
import re
import fcntl
import threading
import time


def find_start(p, start=0):
    """
    Sucht nach dem ersten Vorkommen des Strings "<gpt>" innerhalb des Textes von `p` ab der angegebenen Startposition.

    :param p: Ein Objekt, das den zu durchsuchenden Text im Attribut `text` enthält.
    :param start: Die Startposition innerhalb von `p.text`, ab der die Suche beginnen soll. Standardwert ist 0.
    :return: Die Position des ersten Vorkommens von "<gpt>" in `p.text`. Gibt `-1` zurück, wenn "<gpt>" nicht gefunden wird.
    """
    index = p.text.find("<gpt>", start)
    return index


def find_end(p, start=0):
    """
    Sucht nach dem ersten Vorkommen des Strings "</gpt>" innerhalb des Textes von `p` ab der angegebenen Startposition.

    :param p: Ein Objekt, das den zu durchsuchenden Text im Attribut `text` enthält.
    :param start: Die Startposition innerhalb von `p.text`, ab der die Suche beginnen soll. Standardwert ist 0.
    :return: Die Position des ersten Vorkommens von "</gpt>" in `p.text`. Gibt `-1` zurück, wenn "</gpt>" nicht gefunden wird.
    """
    index = p.text.find("</gpt>", start)
    return index


def collect_messages(docx):
    """
    Durchläuft ein Docx-Dokument und sammelt alle Teilstrings die sich zwischen den Start- und Endtags befinden mit Positionsangaben in einem Array.

    :param docx: Das Docx-Dokument in dem nach GPT-Nachrichten gesucht werden soll.
    :return: Ein Array mit den Textpassagen und deren Startindex, Startabsatz, Endindex, Endabsatz.
    """
    arr = []
    zeile = [-1, -1, -1, -1]
    c = 0

    for i in range(len(docx.paragraphs)):

        p = docx.paragraphs[i]
        c = 0

        while find_start(p, c) > -1 or find_end(p, c) > -1:

            if zeile[0] == -1:
                m = find_start(p, c)
                if m > -1:
                    zeile[0] = m
                    zeile[1] = i
                    c = m + 4

            if zeile[2] == -1:
                n = find_end(p, c)
                if n > -1:
                    zeile[2] = n
                    zeile[3] = i
                    c = n + 5

            if zeile[0] > -1 and zeile[1] > -1 and zeile[2] > -1 and zeile[3] > -1:
                arr.append(
                    [
                        docx_tools.extractTextBetween(
                            zeile[0] + 5, zeile[1], zeile[2], zeile[3], docx
                        ),
                        zeile[0] + 5,
                        zeile[1],
                        zeile[2] - 1,
                        zeile[3],
                    ]
                )
                zeile = [-1, -1, -1, -1]

    return arr


def check_gpt_pairs_and_no_nesting(string):
    """
    Die Funktion prüft ob auf jedes gpt-Opentag ein gpt-Closetag folgt und ob Verschachtelungen vorkommen.

    :param string: Eingabestring, der auf Gültigkeit der gpt-Tags geprüft wird.
    :return: True, wenn jedem Opentag ein Closetag folgt und wenn keine Verschachtelungen vorkommen. Sonst False.
    """
    open_tag = 0
    i = 0
    while i < len(string):
        if string[i : i + 5] == "<gpt>":
            if open_tag > 0:
                return False
            open_tag += 1
            i += 5
        elif string[i : i + 6] == "</gpt>":
            if open_tag == 0:
                return False
            open_tag -= 1
            i += 6
        else:
            i += 1
    return open_tag == 0


def replace_messages_by_answers(messages):
    """
    Ersetzt die Nachrichten für GPT durch die Antworten von GPT.

    :param messages: Array, so wie es von der Funktion "collect_messages" erzeugt wird.
    :return: messages-Array mit anderen Strings
    """
    for i in range(len(messages)):
        messages[i][0] = gpt.send_message(messages[i][0])
    return messages


def insert_answers_into_doc(messages, doc):
    for i in range(len(messages) - 1, -1, -1):
        m = messages[i][1]
        n = messages[i][3]
        p_start = messages[i][2]
        p_end = messages[i][4]
        text = messages[i][0]
        docx_tools.replaceDocTextSegment(m - 5, p_start, n + 6, p_end, doc, text)


def validate_and_process_gpt_tags(cell):
    """
    Überprüft, ob in einem gegebenen Zell-Objekt die '<gpt>' und '</gpt>' Tags korrekt gepaart und nicht verschachtelt sind.
    Verarbeitet dann die Nachrichten innerhalb der Tags, ersetzt sie durch entsprechende Antworten und fügt diese in das Dokument ein.

    Parameter:
    - cell: Das Zell-Objekt, das auf gültige '<gpt>' und '</gpt>' Tags überprüft und verarbeitet wird.

    Return:
    - Ein leerer String, symbolisch für das Fehlen eines expliziten Rückgabewertes, da die Funktion direkt das Dokument modifiziert.
    """

    if not check_gpt_pairs_and_no_nesting(docx_tools.combineDocText(cell)):
        raise Exception(
            "Ungültige Struktur: Für jedes '<gpt>' muss ein korrespondierendes '</gpt>' vorhanden sein. Verschachtelungen sind nicht erlaubt. Bitte überprüfen Sie die Paarung der Tags."
        )
    messages = collect_messages(cell)
    messages = replace_messages_by_answers(messages)
    insert_answers_into_doc(messages, cell)
    return ""


def start_processing(doc):
    validate_and_process_gpt_tags(doc)
    docx_tools.concatTableTexts(doc, validate_and_process_gpt_tags)


def work():

    file_input = config.TOMEDO_CACHE_PROXY
    briefe = glob.glob(file_input + "*.docx")

    for brief in briefe:
        doc = docx.Document(brief)
        pattern = r"\$\[.+\]\$"
        string = docx_tools.combineDocText(doc)
        string += docx_tools.extractInnerDocText(doc)
        if not re.search(pattern, string):
            pattern = r"<gpt>.*\<\/gpt>"
            if re.search(pattern, string, re.DOTALL):
                print("found: " + brief)
                start_processing(doc)
                doc.save(brief)


def loading_animation(event):
    animation = "|/-\\"
    idx = 0
    while not event.is_set():
        print(animation[idx % len(animation)], end="\r")
        idx += 1
        time.sleep(0.1)


def main(worker_function):
    lock_file_path = "/tmp/my_script.lock"

    try:
        lock_file = open(lock_file_path, "w")
        fcntl.flock(lock_file, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except IOError:
        print("Das Skript wird bereits ausgeführt. Beende.")
        sys.exit(1)

    # Ereignis zur Steuerung der Ladeanimation
    stop_event = threading.Event()

    # Starte die Ladeanimation
    animation_thread = threading.Thread(target=loading_animation, args=(stop_event,))
    animation_thread.start()

    try:
        worker_function()
        print("done 👍")
    finally:
        # Stoppe die Ladeanimation
        stop_event.set()
        animation_thread.join()

    fcntl.flock(lock_file, fcntl.LOCK_UN)
    lock_file.close()


if __name__ == "__main__":
    main(work)
